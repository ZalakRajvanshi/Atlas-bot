"""Document ingestion — PDF / DOCX / text into page-tagged chunks.

Chunks carry their source page so Atlas can cite "p. 42" rather than asserting
things about a document the user cannot verify.
"""

from __future__ import annotations

import io
import logging
import re

log = logging.getLogger(__name__)

CHUNK_CHARS = 2400
CHUNK_OVERLAP = 300

# Enough signal to classify a filing without spending a model call.
DOC_TYPE_PATTERNS: list[tuple[str, str]] = [
    (r"\bannual report\b|\bform\s*10-?k\b", "10-K"),
    (r"\bquarterly report\b|\bform\s*10-?q\b", "10-Q"),
    (r"\bform\s*8-?k\b|\bcurrent report\b", "8-K"),
    (r"\bproxy statement\b|\bdef\s*14a\b", "DEF 14A"),
    (r"\bearnings (call|presentation|release)\b|\bq[1-4]\s*20\d\d\b", "earnings"),
    (r"\bprospectus\b|\bform\s*s-?1\b", "S-1"),
    (r"\binvestor (deck|presentation)\b|\bpitch deck\b", "deck"),
    (r"\bequity research\b|\binitiating coverage\b|\bprice target\b", "research"),
]


def _clean(text: str) -> str:
    text = text.replace("\x00", " ")
    # PDF extraction leaves ragged whitespace; collapse it without losing paras.
    text = re.sub(r"[ \t ]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(data: bytes) -> list[tuple[int, str]]:
    """Return [(page_number, text)] — 1-indexed pages."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[tuple[int, str]] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001 — one bad page shouldn't kill ingest
            log.warning("PDF page %s failed to extract: %s", i, exc)
            text = ""
        text = _clean(text)
        if text:
            pages.append((i, text))
    return pages


def extract_docx(data: bytes) -> list[tuple[int, str]]:
    import docx

    document = docx.Document(io.BytesIO(data))
    paras = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    text = _clean("\n".join(paras))
    return [(1, text)] if text else []


def extract_html(data: bytes) -> list[tuple[int, str]]:
    """Strip an HTML document down to readable text.

    SEC EDGAR serves filings as .htm, not PDF, so a user who saves a 10-K
    straight from the source has an HTML file. Rejecting it would be a strange
    gap in a bot that makes a point of citing filings.

    Regex rather than a parser: filing HTML is machine-generated, structurally
    simple, and we only need the text. A parser dependency would earn nothing.
    """
    for encoding in ("utf-8", "latin-1"):
        try:
            markup = data.decode(encoding, errors="ignore")
            break
        except (UnicodeDecodeError, LookupError):
            continue
    else:
        return []

    # SEC inline-XBRL filings open with a large hidden block of machine tags
    # (`false2024FY0000320193P1Y...`). Left in, it becomes the first thing any
    # summariser reads and it defeats document-type detection, so it goes first.
    markup = re.sub(r"<ix:header[^>]*>.*?</ix:header>", " ", markup, flags=re.DOTALL | re.I)
    markup = re.sub(r"<ix:hidden[^>]*>.*?</ix:hidden>", " ", markup, flags=re.DOTALL | re.I)
    markup = re.sub(
        r'<div[^>]*style="[^"]*display:\s*none[^"]*"[^>]*>.*?</div>',
        " ",
        markup,
        flags=re.DOTALL | re.I,
    )

    # Script and style content is never readable text.
    markup = re.sub(
        r"<(script|style)[^>]*>.*?</\1>", " ", markup, flags=re.DOTALL | re.I
    )
    # Block-level tags become line breaks so paragraphs survive.
    markup = re.sub(
        r"</?(p|div|br|tr|h[1-6]|li|table)[^>]*>", "\n", markup, flags=re.I
    )
    # Table cells become spaces, or adjacent columns run together.
    markup = re.sub(r"</?(td|th)[^>]*>", " ", markup, flags=re.I)
    markup = re.sub(r"<[^>]+>", "", markup)

    # Any XBRL identifiers that survive are long unbroken alphanumeric runs
    # with no spaces — never real prose. Drop them rather than let them count
    # toward the text a reader (or the model) has to wade through.
    markup = re.sub(r"\S{60,}", " ", markup)

    import html as _html

    text = _clean(_html.unescape(markup))
    return [(1, text)] if len(text) > 200 else []


def extract_text(data: bytes) -> list[tuple[int, str]]:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = _clean(data.decode(encoding))
            return [(1, text)] if text else []
        except UnicodeDecodeError:
            continue
    return []


def extract(filename: str, data: bytes) -> list[tuple[int, str]]:
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return extract_pdf(data)
    if lowered.endswith((".docx", ".doc")):
        return extract_docx(data)
    if lowered.endswith((".htm", ".html", ".xhtml")):
        return extract_html(data)
    return extract_text(data)


def chunk_pages(pages: list[tuple[int, str]]) -> list[tuple[int, int | None, str]]:
    """Split page text into overlapping chunks.

    Returns [(ordinal, page, text)]. Splits on paragraph boundaries where
    possible so a chunk rarely cuts a sentence in half; overlap preserves
    context for facts that straddle a boundary.
    """
    chunks: list[tuple[int, int | None, str]] = []
    ordinal = 0

    for page_no, text in pages:
        if len(text) <= CHUNK_CHARS:
            chunks.append((ordinal, page_no, text))
            ordinal += 1
            continue

        start = 0
        while start < len(text):
            end = min(start + CHUNK_CHARS, len(text))
            if end < len(text):
                # Prefer a paragraph break, then a sentence end, in the tail.
                window = text[start:end]
                for sep in ("\n\n", ". ", "\n"):
                    cut = window.rfind(sep)
                    if cut > CHUNK_CHARS * 0.5:
                        end = start + cut + len(sep)
                        break
            piece = text[start:end].strip()
            if piece:
                chunks.append((ordinal, page_no, piece))
                ordinal += 1
            if end >= len(text):
                break
            start = max(end - CHUNK_OVERLAP, start + 1)

    return chunks


def guess_doc_type(filename: str, sample: str) -> str | None:
    haystack = f"{filename}\n{sample[:6000]}".lower()
    for pattern, label in DOC_TYPE_PATTERNS:
        if re.search(pattern, haystack):
            return label
    return None


def guess_ticker(sample: str) -> str | None:
    """Look for an explicit exchange:ticker declaration on the cover page.

    Filings and decks nearly always carry one; guessing from prose would be
    worse than returning nothing, so only high-confidence patterns count.
    """
    patterns = [
        r"\((?:NASDAQ|NYSE|NYSE American|CBOE)[:\s]+([A-Z]{1,5})\)",
        r"\b(?:NASDAQ|NYSE)[:\s]+([A-Z]{1,5})\b",
        r"trading symbol[:\s]+([A-Z]{1,5})\b",
        r"ticker symbol[:\s]+([A-Z]{1,5})\b",
    ]
    head = sample[:8000]
    for pattern in patterns:
        match = re.search(pattern, head, re.IGNORECASE)
        if match:
            return match.group(1).upper()
    return None
